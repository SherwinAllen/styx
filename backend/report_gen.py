from pymongo import MongoClient
import gridfs
import pandas as pd
import docx
import re
from collections import Counter
import tempfile
import matplotlib.pyplot as plt
from fastapi.responses import FileResponse
from flask import Flask
import datetime
import os
import json
import hashlib
from bson import Binary

app = Flask(__name__)

# ---------------- MongoDB Setup ----------------
client = MongoClient("mongodb://localhost:27017/")
db = client["forensic_evidence"]
fs = gridfs.GridFS(db)

print("Files in GridFS:")
for filename in fs.list():
    print(filename)

def hash_binary_data(binary_data):
    """Hash binary data from MongoDB"""
    if isinstance(binary_data, Binary):
        # If it's a BSON Binary object
        data_bytes = binary_data
    else:
        # If it's already bytes
        data_bytes = binary_data
    
    hash_object = hashlib.sha256(data_bytes)
    return hash_object.hexdigest()

def get_file_from_mongo(filename):
    """Fetch a file from MongoDB GridFS and return content as text."""
    file_doc = fs.find_one({"filename": filename})
    if not file_doc:
        print(f"[-] File '{filename}' not found in MongoDB.")
        return "", ""
    try:
        print(f"File {filename} Found")
        data = file_doc.read()
        file_hash = hash_binary_data(data)
        # Decode text files; binary files can be handled separately if needed
        return data.decode("utf-8", errors="ignore"), file_hash
    except Exception as e:
        print(f"[!] Error reading {filename} from MongoDB: {e}")
        return "", ""

def extract_logs_from_file(filepath):
    """Reads up to 20 lines from the given file."""
    parsed_data = []
    raw_lines = []
    file_hash = ""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as file:
            content = file.read()
            file_hash = hash_binary_data(content.encode('utf-8'))
            lines = content.splitlines()
            for line in lines[:20]:
                raw_lines.append(line.rstrip('\n'))
                parsed_data.append(line.strip().split())
    except PermissionError as e:
        print(f"Permission denied when accessing {filepath}: {e}")
    return parsed_data, raw_lines, file_hash

def parse_sensor_line(line):
    """Extracts structured info for each sensor log line."""
    pattern = re.compile(
        r'(?P<sensor_id>0x[0-9a-f]+)\).*active-count\s*=\s*(?P<active_count>\d+);'
        r'.*sampling_period\(ms\)\s*=\s*\{(?P<sampling>[\d., ]+)\}.*'
        r'batching_period\(ms\)\s*=\s*\{(?P<batching>[\d., ]+)\}.*selected\s*=\s*(?P<selected>[\d.]+) ms'
    )
    match = pattern.search(line)
    if match:
        return [
            match.group("sensor_id"),
            match.group("active_count"),
            match.group("sampling").replace(",", ", "),
            match.group("batching").replace(",", ", "),
            match.group("selected")
        ]
    else:
        return [""] * 5

def parse_account_info(log_text, file_hash):
    """
    Parses Android account-related forensic dump text into two DataFrames:
      - Accounts table
      - Registered Services table
    """
    accounts = []
    services = []

    # --- Parse accounts ---
    for line in log_text.splitlines():
        line = line.strip()
        # Match Account {name=..., type=...}
        acc_match = re.match(r'Account\s*\{name=([^,]+),\s*type=([^}]+)\}', line)
        if acc_match:
            name = acc_match.group(1).strip()
            acc_type = acc_match.group(2).strip()
            accounts.append({"Account Name": name, "Type": acc_type})
            continue

        # Match ServiceInfo lines
        svc_match = re.match(
            r'ServiceInfo:\s*AuthenticatorDescription\s*\{type=([^}]+)\},\s*ComponentInfo\{([^}]+)\},\s*uid\s*(\d+)',
            line
        )
        if svc_match:
            svc_type = svc_match.group(1).strip()
            component = svc_match.group(2).strip()
            uid = svc_match.group(3).strip()
            services.append({
                "Type": svc_type,
                "Component": component,
                "UID": uid
            })

    # Convert to DataFrames
    accounts_df = pd.DataFrame(accounts)
    services_df = pd.DataFrame(services)

    return accounts_df, services_df, file_hash

def extract_sensor_timestamps(raw_lines, file_hash):
    """
    Extract timestamps, sensor IDs, and (optional) sensor names.
    Returns dict: {sensor_id: DataFrame}
    """
    sensor_data = {}
    # Regex captures timestamps, sensor IDs, and possible sensor names
    pattern = re.compile(
        r'(?P<timestamp>\d{2}-\d{2}\s\d{2}:\d{2}:\d{2}\.\d+).*?'
        r'(?P<sensor_id>0x[0-9a-f]+).*?'
        r'(?:SensorName\s*=\s*(?P<name>[A-Za-z0-9_ -]+))?',
        re.IGNORECASE
    )

    for line in raw_lines:
        match = pattern.search(line)
        print(line)
        if match:
            timestamp = match.group("timestamp")
            sensor_id = match.group("sensor_id")
            sensor_name = match.group("name") or "Unknown Sensor"
            if sensor_id not in sensor_data:
                sensor_data[sensor_id] = {"name": sensor_name, "entries": []}
            sensor_data[sensor_id]["entries"].append((timestamp, line.strip()))

    # Convert to DataFrames
    sensor_dfs = {}
    for sensor_id, info in sensor_data.items():
        df = pd.DataFrame(info["entries"], columns=["Timestamp", "Log Line"])
        sensor_dfs[sensor_id] = (info["name"], df)

    return sensor_dfs, file_hash

def add_dataframe_to_doc(doc, df, title, max_cols_per_table=5):
    """
    Writes a pandas DataFrame into the Word doc as formatted tables.
    Splits wide DataFrames into multiple tables if columns exceed max_cols_per_table.
    """
    if df.empty:
        doc.add_paragraph(f"{title} - No data found.\n", style='Heading3')
        return

    columns = df.columns.tolist()
    start = 0
    table_index = 1

    while start < len(columns):
        subset_cols = columns[start:start + max_cols_per_table]
        sub_df = df[subset_cols]

        # Add title for this part
        if len(columns) > max_cols_per_table:
            doc.add_paragraph(f"{title} (Part {table_index})", style='Heading4')
        else:
            doc.add_paragraph(title, style='Heading3')

        # Create table
        table = doc.add_table(rows=1, cols=len(sub_df.columns))
        table.style = "Table Grid"

        # Header row
        for i, col_name in enumerate(sub_df.columns):
            table.cell(0, i).text = col_name
            for run in table.cell(0, i).paragraphs[0].runs:
                run.bold = True

        # Data rows
        for _, row in sub_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        doc.add_paragraph("\n")
        start += max_cols_per_table
        table_index += 1

def extract_sensor_data(log_text, file_hash):
    sensors = {}
    current_sensor = None
    records = []

    for line in log_text.splitlines():
        line = line.strip()

        # --- Detect start of a new sensor section ---
        match_sensor = re.match(r'^(.*?):.*events$', line)
        print(line)
        if match_sensor:
            # Save previous sensor's records
            if current_sensor and records:
                sensors[current_sensor] = pd.DataFrame(records)
                records = []
            current_sensor = match_sensor.group(1).strip()
            continue

        # --- Match event lines ---
        # Flexible regex for formats like:
        # 1 (ts=123.456, wall=12:34:56.789) 1.00, 0.00,
        match_event = re.match(
            r'^\d+\s*\(ts=([\d.]+),\s*wall=([\d:.]+)\)\s*(.*)', line
        )
        if match_event:
            ts = float(match_event.group(1))
            wall = match_event.group(2).strip()
            values_str = match_event.group(3).strip()

            # --- Handle value formats ---
            if "[value masked]" in values_str:
                record = {"ts": ts, "wall_time": wall, "values": "[value masked]"}
            elif values_str:
                # Extract all numeric values, even if followed by commas
                nums = [float(x) for x in re.findall(r'[-+]?\d*\.\d+|\d+', values_str)]
                record = {"ts": ts, "wall_time": wall}
                for i, val in enumerate(nums, start=1):
                    record[f"value_{i}"] = val
            else:
                record = {"ts": ts, "wall_time": wall}

            records.append(record)

    # Save last sensor's data
    if current_sensor and records:
        sensors[current_sensor] = pd.DataFrame(records)

    return sensors, file_hash

def parse_bluetooth_log(doc, text, file_hash):
    """Extracts Bluetooth connection and bonded device info."""

    # 1⃣ Connection / Disconnection events
    print(text)
    dates = re.findall(r"(\d{2}-\d{2})\s\d{2}:\d{2}:\d{2}\.\d{3}", text)
    print(dates)

    # Count number of events per day
    counter = Counter(dates)

    # Convert to pandas Series for plotting
    df = pd.Series(counter).sort_index()

    # --- 📊 Calculate statistics ---
    avg_events = df.mean() if not df.empty else 0
    above_avg_days = (df > avg_events).sum()
    below_avg_days = (df < avg_events).sum()
    total_days = len(df)

    perc_above = (above_avg_days / total_days * 100) if total_days else 0
    perc_below = (below_avg_days / total_days * 100) if total_days else 0

    # --- 📈 Plot bar graph ---
    title = "Bluetooth Events Per Day"
    plt.figure(figsize=(10, 6))
    df.plot(kind='bar', color='skyblue')
    plt.xlabel("Date (MM - DD)")
    plt.ylabel("Number of Bluetooth Events")    
    plt.title(title)
    plt.xticks(rotation=45)

    # Save plot to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
        plt.savefig(tmpfile.name, format='png')
        plt.close()
        doc.add_paragraph(title, style='Heading3')
        doc.add_picture(tmpfile.name, width=docx.shared.Inches(6))

    # --- 📄 Add statistics table ---
    doc.add_paragraph("Bluetooth Activity Statistics:", style='Heading4')

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light List Accent 1'

    # Fill the table
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Average Bluetooth Events per Day"
    table.cell(1, 1).text = f"{avg_events:.2f}"
    table.cell(2, 0).text = "Days Above Average"
    table.cell(2, 1).text = f"{above_avg_days} ({perc_above:.1f}%)"
    table.cell(3, 0).text = "Days Below Average"
    table.cell(3, 1).text = f"{below_avg_days} ({perc_below:.1f}%)"
    table.cell(4, 0).text = "Total Days Analyzed"
    table.cell(4, 1).text = f"{total_days}"

    # Adjust column widths (optional, looks cleaner)
    for row in table.rows:
        row.cells[0].width = docx.shared.Inches(3)
        row.cells[1].width = docx.shared.Inches(2)

    doc.add_paragraph()  # spacing after the table

    # 2⃣ Bonded devices
    bonded_pattern = r"\s*\(Connected\)\s*([0-9A-F:]{17}) \[.*?\] ([^\(]+)"
    bonded_match = re.findall(bonded_pattern, text)

    bonded_devices = []
    if bonded_match:
        for match in bonded_match:
            bonded_devices.append({
                "Device Name": match[1].strip(),
                "MAC Address": match[0]
            })

    df_bonded = pd.DataFrame(bonded_devices)

    return df_bonded, file_hash

def extract_ip_info(output_text, file_hash):
    """
    Parse ADB `ip addr` output and return a pandas DataFrame
    with columns: Interface, Status, MTU, IPv4, Broadcast, IPv6, MAC, Notes
    mapped to value_1, value_2, ..., value_8 for consistency.
    """
    interfaces = []
    current_iface = {}
    for line in output_text.splitlines():
        line = line.strip()
        if re.match(r'^\d+:', line):
            # New interface line
            if current_iface:
                interfaces.append(current_iface)
            m = re.match(r'^(\d+):\s+([\w@]+):\s+<([^>]*)>.*mtu\s+(\d+)', line)
            if m:
                current_iface = {
                    "Interface": m.group(2),
                    "Status": "UP" if "UP" in m.group(3).split(",") else "DOWN",
                    "MTU": int(m.group(4)),
                    "IPv4": "",
                    "Broadcast": "",
                    "IPv6": "",
                    "MAC": "",
                    "Notes": ""
                }
        elif line.startswith("link/"):
            # MAC address line
            m = re.match(r'link/\w+\s+([\da-f:]+)', line)
            if m:
                current_iface["MAC"] = m.group(1)
        elif line.startswith("inet "):
            # IPv4 line
            m = re.match(r'inet\s+([\d./]+)\s+brd\s+([\d.]+)', line)
            if m:
                current_iface["IPv4"] = m.group(1)
                current_iface["Broadcast"] = m.group(2)
        elif line.startswith("inet6 "):
            # IPv6 line
            m = re.match(r'inet6\s+([\da-f:]+/[\d]+)', line)
            if m:
                current_iface["IPv6"] = m.group(1)

    if current_iface:
        interfaces.append(current_iface)

    # Convert to DataFrame and map to value_1..value_8
    df = pd.DataFrame(interfaces)
    
    return df, file_hash

def get_location(location_file):
    with open(location_file,'r', encoding = 'utf-8') as f:
        location_text = f.read()
        file_hash = hash_binary_data(location_text.encode('utf-8'))
        regex = re.compile(
            r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?(?:hAcc=(?P<acc>\d+\.?\d*))?',
            re.IGNORECASE | re.DOTALL
        )

        matches = list(regex.finditer(location_text))
        records = []
        for m in matches:
            records.append({
                "provider": m.group("provider"),
                "latitude": float(m.group("lat")),
                "longitude": float(m.group("lon")),
                "accuracy": float(m.group("acc")) if m.group("acc") else None
            })

        # Convert to DataFrame
        df = pd.DataFrame(records)

    return df, file_hash

def parse_wifi_log_extended(log_text: str, file_hash: str):
    """
    Parse ADB Wi-Fi diagnostic logs including:
      - SSID/BSSID connection info
      - Connection metrics
      - Supplicant state transitions
      - Multi-Link (Mlink) info
    Returns a dict of DataFrames.
    """
    dfs = {}

    # 1⃣ ---- Wi-Fi SSID/BSSID Info ----
    ssid_pattern = re.compile(
        r'rec\[\d+\]:\s+'                      # rec number
        r'time=(?P<timestamp>[\d\-:\. ]+)\s+'  # timestamp
        r'processed=(?P<processed>\S+)\s+'
        r'org=(?P<org>\S+)\s+'
        r'dest=(?P<dest>\S+)\s+'
        r'what=(?P<what>\S+)\s+'
        r'screen=\S+\s+\d+\s+\d+\s+'
        r'ssid:\s*"(?P<ssid>[^"]+)"\s+'
        r'bssid:\s*(?P<bssid>[0-9a-f:]+)\s+'
        r'nid:\s*(?P<nid>\d+)\s+'
        r'frequencyMhz:\s*(?P<freq>\d+)\s+'
        r'state:\s*COMPLETED', 
        re.IGNORECASE
    )
    ssid_records = []
    for line in log_text.splitlines():
        m = ssid_pattern.search(line)
        if m:
            ssid_records.append({
                "timestamp": m.group('timestamp'),
                "ssid": m.group('ssid'),
                "bssid": m.group('bssid')
            })

    # 2⃣ ---- Wi-Fi Metrics ----
    wifi_pattern = re.compile(
        r"time=(?P<time>[\d\-\s:]+).*?"
        r"session=(?P<session>[^,]+),?"
        r".*?netid=(?P<netid>[^,]+),?"
        r".*?rssi=(?P<rssi>[^,]+),?"
        r".*?filtered_rssi=(?P<filtered_rssi>[^,]+),?"
        r".*?freq=(?P<freq>[^,]+),?"
        r".*?txLinkSpeed=(?P<txLinkSpeed>[^,]+),?"
        r".*?rxLinkSpeed=(?P<rxLinkSpeed>[^,]+),?",
        re.DOTALL
    )
    wifi_records = []
    for line in log_text.splitlines():
        if "rssi=" in line and "txLinkSpeed=" in line:
            m = wifi_pattern.search(line)
            if m:
                wifi_records.append(m.groupdict())

    # 3⃣ ---- Supplicant State Tracker ----
    supplicant_pattern = re.compile(
        r"rec\[\d+\]: time=(?P<time>[\d\-:\.\s]+).*?"
        r"org=(?P<org_state>\S+).*?"
        r"dest=(?P<dest_state>\S*).*?"  # allow empty dest
        r"what=(?P<what>[0-9xXA-F]+)",
        re.DOTALL
    )
    supplicant_records = [m.groupdict() for m in supplicant_pattern.finditer(log_text)]

    # Filter out rows where dest_state is empty
    supplicant_records = [r for r in supplicant_records if r["dest_state"].strip() != "<null>"]

    # Convert to DataFrame
    if supplicant_records:
        dfs["supplicant_states"] = pd.DataFrame(supplicant_records)

    # 4⃣ ---- Mlink / Multi-Link Operation ----
    mlink_pattern = re.compile(
        r"\{linkId=(?P<linkId>\d+),linkRssi=(?P<linkRssi>[^,]+),linkFreq=(?P<linkFreq>[^,]+),"
        r"txLinkSpeed=(?P<txLinkSpeed>[^,]+),rxLinkSpeed=(?P<rxLinkSpeed>[^,]+).*?\}",
        re.DOTALL
    )

    mlink_records = [m.groupdict() for m in mlink_pattern.finditer(log_text)]

    # Convert to DataFrames
    if ssid_records:
        dfs["wifi_networks"] = pd.DataFrame(ssid_records)
    if wifi_records:
        dfs["wifi_metrics"] = pd.DataFrame(wifi_records)
    if supplicant_records:
        dfs["supplicant_states"] = pd.DataFrame(supplicant_records)
    if mlink_records:
        dfs["mlink_info"] = pd.DataFrame(mlink_records)

    return dfs, file_hash

def parse_location_data(loc_path):
    with open(loc_path, 'r', encoding="utf-8") as f:
        log_text = f.read()
        file_hash = hash_binary_data(log_text.encode('utf-8'))
    
    regex = re.compile(
        r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*'
        r'(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?'
        r'(?:hAcc=(?P<acc>\d+\.?\d*))?',
        re.IGNORECASE | re.DOTALL
    )

    matches = list(regex.finditer(log_text))
    if not matches:
        print("[-] No coordinate patterns found in dumpsys output.")
        df = pd.DataFrame(columns=["timestamp", "provider", "lat", "lon", "accuracy"])
    else:
        rows = []
        for m in matches:
            provider = m.group('provider') or 'unknown'
            lat = float(m.group('lat'))
            lon = float(m.group('lon'))
            acc = m.group('acc')
            ts = datetime.datetime.now().isoformat()
            rows.append({
                "timestamp": ts,
                "provider": provider,
                "lat": lat,
                "lon": lon,
                "accuracy": float(acc) if acc else None
            })

        df = pd.DataFrame(rows)
        print(f" Parsed {len(df)} location entries.")
        print(df)

    return df, file_hash

def parse_trust_manager_states(log_text, file_hash):
    """
    Parse Android Trust Manager states from forensic dump text.
    Returns DataFrame with trust state information.
    """
    trust_states = []
    
    # Regex pattern to match trust manager state lines
    pattern = re.compile(
        r'User\s+"(?P<user_name>[^"]+)"\s+'
        r'\(id=(?P<user_id>\d+),\s*flags=(?P<flags>0x[0-9a-fA-F]+)\)\s*'
        r'\(current\):\s*'
        r'trustState=(?P<trust_state>\w+),\s*'
        r'trustManaged=(?P<trust_managed>\d+),\s*'
        r'deviceLocked=(?P<device_locked>\d+),\s*'
        r'isActiveUnlockRunning=(?P<active_unlock_running>\d+),\s*'
        r'strongAuthRequired=(?P<strong_auth_required>0x[0-9a-fA-F]+)',
        re.IGNORECASE
    )
    
    for line in log_text.splitlines():
        line = line.strip()
        match = pattern.search(line)
        if match:
            trust_states.append({
                "User Name": match.group("user_name"),
                "User ID": int(match.group("user_id")),
                "Flags": match.group("flags"),
                "Trust State": match.group("trust_state"),
                "Trust Managed": bool(int(match.group("trust_managed"))),
                "Device Locked": bool(int(match.group("device_locked"))),
                "Active Unlock Running": bool(int(match.group("active_unlock_running"))),
                "Strong Auth Required": match.group("strong_auth_required")
            })
    
    # Convert to DataFrame
    df = pd.DataFrame(trust_states)
    
    return df, file_hash


log_files = {
    "Account Information": "account_information.txt",
    "Bluetooth Information": "bluetooth_information.txt",
    "Device Properties": "device_properties.txt",
    "Sensor Data": "sensor_data.txt",
    "Ip information": "ip_address_information.txt"
}

column_headers = {
    "Account Information": ["Field", "Value"],
    "Bluetooth Information": ["Field", "Value"],
    "Device Properties": ["Field", "Value"],
    "Sensor Data": ["Timestamp", "Sensor Type", "Value"]
}

# ----------------------------------------------------------------
# Function that builds the forensic report and saves it
# ----------------------------------------------------------------
log_files = {
    "Basic Device Properties": "basic_device_info.txt",
    "Account Information": "account_information.txt",
    "Bluetooth Information": "bluetooth_information.txt",
    "Device Properties": "device_properties.txt",
    "Sensor Data": "sensor_data.txt",
    "Ip information": "ip_address_information.txt",
    "WiFi Information": "wifi_information.txt",
    "Location Information": "dumpsys_location.txt",
    "Trust Manager": "trust_information.txt",
    "Notification Information": "notification_information.txt",
    "Keystore Information": "keystore_information.txt"
}

# ---------------- Forensic Report Generation ----------------
def generate_forensic_report(output_dir="downloads"):
    """Generates the forensic .docx report using MongoDB data."""
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, "Preliminary_Forensic_Report.docx")

    doc = docx.Document()
    doc.add_paragraph("Preliminary Forensic Report", style='Title')
    
    # Initialize hash collection
    all_hashes = []
    
    # ---- Basic device properties --------
    basic_prop_text, basic_prop_hash = get_file_from_mongo(log_files["Basic Device Properties"])
    if basic_prop_text.strip():
        # Split lines like "Key: Value" into a 2-column DataFrame
        basic_props = []
        for line in basic_prop_text.splitlines():
            line = line.strip()
            if not line:
                continue
            if ':' in line:
                key, value = map(str.strip, line.split(':', 1))
                basic_props.append({"Property": key, "Value": value})
            else:
                # Handle cases without colon
                basic_props.append({"Property": line, "Value": ""})
        basic_props_df = pd.DataFrame(basic_props)
        add_dataframe_to_doc(doc, basic_props_df, "Basic Device Properties")
        all_hashes.append({"File": "basic_device_info.txt", "SHA256 Hash": basic_prop_hash})
    else:
        doc.add_paragraph("Basic Device Properties - No data found.\n", style='Heading3')

    # --- Account Info ---
    acc_text, acc_hash = get_file_from_mongo(log_files["Account Information"])
    acc_df, service_df, acc_hash = parse_account_info(acc_text, acc_hash)
    add_dataframe_to_doc(doc, acc_df, "Account Information")
    add_dataframe_to_doc(doc, service_df, "Service Information")
    all_hashes.append({"File": "account_information.txt", "SHA256 Hash": acc_hash})

    # --- Wi-Fi Info ---
    wifi_text, wifi_hash = get_file_from_mongo(log_files["WiFi Information"])
    wifi_df_dict, wifi_hash = parse_wifi_log_extended(wifi_text, wifi_hash)
    for section_name, df in wifi_df_dict.items():
        add_dataframe_to_doc(doc, df, f"Wi-Fi: {section_name.replace('_', ' ').title()}")
    all_hashes.append({"File": "wifi_information.txt", "SHA256 Hash": wifi_hash})

    # --- Bluetooth Info ---
    bt_text, bt_hash = get_file_from_mongo(log_files["Bluetooth Information"])
    df_bonded, bt_hash = parse_bluetooth_log(doc, bt_text, bt_hash)
    add_dataframe_to_doc(doc, df_bonded, "Bonded Bluetooth Devices")
    all_hashes.append({"File": "bluetooth_information.txt", "SHA256 Hash": bt_hash})

    # --- Location Info ---
    loc_text, loc_hash = get_file_from_mongo(log_files["Location Information"])
    loc_df, loc_hash = get_location_text(loc_text, loc_hash)
    add_dataframe_to_doc(doc, loc_df, "Location Information")
    all_hashes.append({"File": "dumpsys_location.txt", "SHA256 Hash": loc_hash})

    # --- Sensor Data ---
    sensor_text, sensor_hash = get_file_from_mongo(log_files["Sensor Data"])
    sensor_dataframes, sensor_hash = extract_sensor_data(sensor_text, sensor_hash)
    for sensor_name, df in sensor_dataframes.items():
        add_dataframe_to_doc(doc, df, sensor_name)
    all_hashes.append({"File": "sensor_data.txt", "SHA256 Hash": sensor_hash})

    # --- IP Info ---
    ip_text, ip_hash = get_file_from_mongo(log_files["Ip information"])
    ip_df, ip_hash = extract_ip_info(ip_text, ip_hash)
    add_dataframe_to_doc(doc, ip_df, "IP Address Information")
    all_hashes.append({"File": "ip_address_information.txt", "SHA256 Hash": ip_hash})

    # --- Trust Manager Information ---

    trust_text, trust_hash = get_file_from_mongo(log_files["Trust Manager"])
    trust_df, trust_hash = parse_trust_manager_states(trust_text, trust_hash)
    add_dataframe_to_doc(doc, trust_df, "Trust Manager State Information")
    all_hashes.append({"File": "trust_manager_states.txt", "SHA256 Hash": trust_hash})
    
    # --- adding hashing not summarized ---
    keystore_text, keystore_hash = get_file_from_mongo(log_files["Keystore Information"])
    all_hashes.append(({"File": "keystore_information.txt", "SHA256 Hash": keystore_hash}))
    
    notification_text, notification_hash = get_file_from_mongo(log_files["Notification Information"])
    all_hashes.append(({"File": "notification_information.txt", "SHA256 Hash": notification_hash}))

    # --- Add all hashes in one table at the end ---
    doc.add_paragraph("File Integrity Information", style='Heading1')
    doc.add_paragraph("SHA256 Hashes of All Analyzed Files", style='Heading2')

    
    if all_hashes:
        # Create combined hash DataFrame
        hash_df = pd.DataFrame(all_hashes)
        hash_df["Analysis Date"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        add_dataframe_to_doc(doc, hash_df, "File Integrity Hashes")
    else:
        doc.add_paragraph("No file integrity information available.", style='Heading3')

    # Save to DOCX
    doc.save(output_path)
    print(f"Forensic report saved to: {output_path}")
    return output_path

# ---------------- Flask Route ----------------
@app.get("/download_report")
def download_report():
    report_path = "downloads/Preliminary_Forensic_Report.docx"
    if os.path.exists(report_path):
        return FileResponse(
            path=report_path,
            filename="Preliminary_Forensic_report.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return {"error": "Report not found"}


# ---------------- Helper for location text ----------------
def get_location_text(location_text, file_hash):
    """Parse location text from MongoDB (previously from file)."""
    regex = re.compile(
        r'Location\[(?:provider=)?(?P<provider>[\w\-]+)?\s*'
        r'(?P<lat>-?\d+\.\d+)[, ]+(?P<lon>-?\d+\.\d+).*?'
        r'(?:hAcc=(?P<acc>\d+\.?\d*))?',
        re.IGNORECASE | re.DOTALL
    )
    matches = list(regex.finditer(location_text))
    records = []
    for m in matches:
        records.append({
            "provider": m.group("provider") or "unknown",
            "latitude": float(m.group("lat")),
            "longitude": float(m.group("lon")),
            "accuracy": float(m.group("acc")) if m.group("acc") else None
        })
    
    df = pd.DataFrame(records)
    
    return df, file_hash

# ---------------- Main ----------------
if __name__ == "__main__":
    generate_forensic_report()